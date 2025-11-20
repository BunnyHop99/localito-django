from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Función para sombrear celdas
def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# Crear documento
doc = Document()

# Título principal
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("[LOGO DE LA EMPRESA]\nCOTIZACIÓN ESPECIAL – PROYECTO PILOTO")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 70, 122)

doc.add_paragraph("\n")

# Datos generales
doc.add_heading("Información del Cliente", level=2)
doc.add_paragraph(
    "Cliente: [Nombre del Cliente]\n"
    "Empresa: [Razón Social]\n"
    "RFC: [RFC]\n"
    "Fecha: 3/11/2025\n"
    "Válida por 30 días"
)

doc.add_heading("Emitido por", level=2)
doc.add_paragraph(
    "[Nombre de tu Empresa]\n[Dirección completa] | [Teléfono] | [Email] | [Sitio web]"
)

doc.add_paragraph("\n")

# Resumen ejecutivo con color
summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = summary.add_run(
    "Resumen Ejecutivo:\n"
    "Le presentamos esta propuesta especial como primer cliente piloto del sistema. "
    "El Sistema Integral de Gestión Empresarial permite administrar ventas, inventarios, "
    "facturación CFDI y reportes desde una sola plataforma moderna, segura y fácil de usar.\n"
    "Tiempo de entrega: 4 semanas\nInversión especial: $98,000 MXN + IVA"
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(50, 50, 50)

doc.add_paragraph("\n")

# Lista de módulos con estilo
modules = [
    ("Punto de Venta (POS)", 48, "$25,000", [
        ("Diseño UI/UX", 12, "Pantallas amigables y flujo de ventas rápido."),
        ("Desarrollo flujo de ventas y caja", 20, "Registro de ventas, descuentos, métodos de pago."),
        ("Integración CFDI", 8, "Generación automática de facturas electrónicas."),
        ("Pruebas y ajustes", 8, "Simulación de ventas y corrección de errores.")
    ]),
    ("Control de Inventario", 40, "$20,000", [
        ("Diseño base de datos", 10, "Registro de productos, lotes y movimientos."),
        ("Desarrollo funcionalidades", 20, "Entradas y salidas de productos, alertas de stock."),
        ("Pruebas y validación", 10, "Simulación para garantizar consistencia de datos.")
    ]),
    ("Facturación CFDI 4.0", 40, "$20,000", [
        ("Integración PAC", 15, "Conexión segura con Facturama/Finkok."),
        ("Generación y validación", 15, "Emisión de CFDI con validación de RFC e impuestos."),
        ("Pruebas de facturación", 10, "Simulación de distintos escenarios de facturación.")
    ]),
    ("Dashboard y Reportes", 32, "$18,000", [
        ("Diseño de dashboard y KPIs", 10, "Panel visual con métricas de ventas e inventario."),
        ("Programación de reportes", 15, "Exportación a PDF/Excel con filtros por fecha/producto."),
        ("Pruebas y optimización", 7, "Verificación de consistencia y velocidad de carga.")
    ]),
    ("Usuarios, Configuración y Capacitación", 32, "$15,000", [
        ("Gestión de roles y permisos", 12, "Niveles de acceso: administrador, vendedor, gerente."),
        ("Configuración inicial", 10, "Datos fiscales, catálogo, impuestos, reglas de negocio."),
        ("Capacitación y documentación", 10, "Manual de usuario y sesión práctica para el personal.")
    ])
]

doc.add_heading("Detalle de Módulos", level=2)

for mod in modules:
    doc.add_heading(mod[0], level=3)
    doc.add_paragraph(f"Duración estimada: {mod[1]} horas | Costo: {mod[2]}")
    doc.add_paragraph("Fases y actividades técnicas:")
    for phase in mod[3]:
        p = doc.add_paragraph(f"- {phase[0]} ({phase[1]} hrs): {phase[2]}")
        p.style = 'List Bullet'

doc.add_paragraph("\n")

# Tabla resumen de inversión con colores
doc.add_heading("Resumen de Inversión", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Módulo"
hdr_cells[1].text = "Horas"
hdr_cells[2].text = "Costo"
for cell in hdr_cells:
    set_cell_background(cell, "D9E1F2")  # color azul claro
    cell.paragraphs[0].runs[0].font.bold = True

for mod in modules:
    row_cells = table.add_row().cells
    row_cells[0].text = mod[0]
    row_cells[1].text = str(mod[1])
    row_cells[2].text = mod[2]

# Fila total
row_total = table.add_row().cells
row_total[0].text = "Total del Proyecto (4 semanas)"
row_total[1].text = str(sum([m[1] for m in modules]))
row_total[2].text = "$98,000 + IVA"
for cell in row_total:
    set_cell_background(cell, "FCE4D6")  # color naranja claro
    cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph("\n")

# Esquema de pagos con color
doc.add_heading("Esquema de Pagos", level=2)
pay_table = doc.add_table(rows=1, cols=3)
hdr = pay_table.rows[0].cells
hdr[0].text = "Etapa"
hdr[1].text = "%"
hdr[2].text = "Monto"
for cell in hdr:
    set_cell_background(cell, "D9E1F2")
    cell.paragraphs[0].runs[0].font.bold = True

pays = [
    ("Anticipo al firmar contrato", "40%", "$39,200"),
    ("Entrega de Módulos POS + Inventario", "30%", "$29,400"),
    ("Entrega final del sistema completo", "30%", "$29,400"),
]

for pay in pays:
    row = pay_table.add_row().cells
    row[0].text, row[1].text, row[2].text = pay

doc.add_paragraph("\n")

# Condiciones especiales
doc.add_heading("Condiciones Especiales – Proyecto Piloto", level=2)
doc.add_paragraph(
    "- Entrega total en 4 semanas.\n"
    "- Soporte técnico premium gratuito durante 6 meses.\n"
    "- Corrección de errores y mejoras sin costo durante el soporte.\n"
    "- Capacitación completa para el personal designado.\n"
    "- Entrega de código fuente, documentación y manual de usuario.\n"
    "- Pagos mediante transferencia bancaria con factura fiscal.\n"
    "- Cliente será reconocido como Primer Caso de Éxito, con beneficios en futuras actualizaciones."
)

# Beneficios adicionales
doc.add_heading("Beneficios Adicionales", level=2)
doc.add_paragraph(
    "✅ Precio preferencial exclusivo (descuento especial).\n"
    "✅ Soporte extendido y atención prioritaria.\n"
    "✅ Actualizaciones sin costo durante el piloto.\n"
    "✅ Participación directa en la mejora del sistema.\n"
    "✅ Garantía total de satisfacción."
)

# Firmas
doc.add_heading("Firmas de Aceptación", level=2)
sign_table = doc.add_table(rows=2, cols=2)
sign_table.cell(0, 0).text = "Representante del Cliente"
sign_table.cell(0, 1).text = "Representante de [Tu Empresa]"
sign_table.cell(1, 0).text = "Nombre: _______________________\nFirma: _______________________\nFecha: _______________________"
sign_table.cell(1, 1).text = "Nombre: _______________________\nFirma: _______________________\nFecha: _______________________"

# Guardar documento
doc.save("Cotizacion_Proyecto_Piloto_Profesional.docx")
print("Archivo generado: Cotizacion_Proyecto_Piloto_Profesional.docx")
